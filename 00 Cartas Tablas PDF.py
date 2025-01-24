import shutil
from docx import Document
import openpyxl
import pandas as pd
from num2words import num2words
import re
import os
import sys
#import win32com.client
#import win32com.client as win32


script_directory = os.path.dirname(os.path.abspath(__file__))
working_folder = os.path.abspath(os.path.join(script_directory, '..'))
function_library = os.path.abspath(os.path.join(script_directory, 'Library'))
sys.path.append(function_library)  # Add the library folder to the path.


#cache_dir = os.path.join(os.environ['LOCALAPPDATA'], 'Temp', 'gen_py')
#if os.path.exists(cache_dir):
#    shutil.rmtree(cache_dir)
#    print("win32com cache cleared.")
#else:
#    print("win32com cache folder not found.")

# Mostrar tablas, llenar tablas, convertir a PDF
def normalize_header(header):
    # Remove non-printable characters but keep #, %, -, and .
    return re.sub(r'[^\w\sÁÉÍÓÚáéíóúÑñ#%()\-.]', '', header).strip()


def populate_table(word_document, df_source, table_number, table_headers_row, table_headers, table_types, table_total, table_df_mapping, table_total_headers):
    # Validate table_types mapping
    for col in table_df_mapping:
        if col not in table_types:
            print(f"[ERROR] Column {col} is missing in table_types!")
        else:
            print(f"[DEBUG] Column: {col}, Type: {table_types[col]}")
    df_source['Precio Unitario'] = pd.to_numeric(df_source['Precio Unitario'], errors='coerce')
    df_source['Importe total Máximo'] = pd.to_numeric(df_source['Importe total Máximo'], errors='coerce')
    df_source['Cantidad Máxima'] = pd.to_numeric(df_source['Cantidad Máxima'], errors='coerce')
    
    # Use the passed-in document
    doc = word_document

    # Check if the table exists
    if table_number > len(doc.tables) or table_number < 1:
        raise ValueError(f"Table number {table_number} does not exist in the Word file.")
    
    # Get the specified table
    table = doc.tables[table_number - 1]

    # Normalize and validate headers
    raw_word_headers = [cell.text for cell in table.rows[table_headers_row - 1].cells]
    normalized_word_headers = [normalize_header(header) for header in raw_word_headers]

    # Debugging to identify which table is problematic
    print(f"[DEBUG] Table Number: {table_number}")
    print(f"[DEBUG] Raw Word Headers: {raw_word_headers}")
    print(f"[DEBUG] Normalized Word Headers: {normalized_word_headers}")
    print(f"[DEBUG] Expected Headers: {table_headers}")

    missing_headers = [header for header in table_headers if header not in normalized_word_headers]
    if missing_headers:
        raise ValueError(f"Missing headers in Word table {table_number}: {missing_headers}")

    # Check if all mappings exist in the DataFrame
    missing_columns = [col for col in table_df_mapping if col not in df_source.columns]
    if missing_columns:
        raise ValueError(f"Missing columns in DataFrame: {missing_columns}")

    # Remove all rows after the header row
    for _ in range(len(table.rows) - table_headers_row):
        table._element.remove(table.rows[table_headers_row]._element)

    # Apply type transformations
    def apply_type_conversion(value, column):
        if pd.isnull(value) or value is None:
            print(f"[DEBUG] Column: {column}, Value is null or None.")
            return ""

        if column in table_types:
            try:
                if table_types[column] == 'currency':
                    value = float(re.sub(r"[^\d.]", "", str(value)))  # Remove non-numeric characters
                    transformed = f"${value:,.2f}"  # Format as currency
                    print(f"[DEBUG] Column: {column}, Raw Value: {value}, Transformed Value: {transformed}")
                    return transformed
                elif table_types[column] == 'number':
                    value = float(re.sub(r"[^\d.]", "", str(value)))  # Remove non-numeric characters
                    transformed = f"{int(value):,}" if value.is_integer() else f"{value:,.2f}"  # Format as number
                    print(f"[DEBUG] Column: {column}, Raw Value: {value}, Transformed Value: {transformed}")
                    return transformed
                elif table_types[column] == 'string':
                    transformed = str(value).strip()
                    print(f"[DEBUG] Column: {column}, Transformed Value: {transformed}")
                    return transformed
            except ValueError as e:
                print(f"[ERROR] Column: {column}, Value: {value}, Error: {e}")
                return str(value).strip()
        print(f"[DEBUG] Column: {column}, No matching type, returning raw value.")
        return str(value).strip()

    # Populate the table
    for index, row in df_source.iterrows():
        new_row = table.add_row().cells
        for col_index, column in enumerate(table_df_mapping):
            value = row[column]
            transformed_value = apply_type_conversion(value, column)  # No need to pass table_types
            print(f"Column: {column}, Raw Value: {value}, Transformed Value: {transformed_value}")
            new_row[col_index].text = transformed_value


    # Handle totals
    if table_total:
        totals = {}
        for total_header in table_total_headers:
            if total_header not in table_headers:
                raise ValueError(f"Total header '{total_header}' is not in Word table headers.")
            header_index = table_headers.index(total_header)
            df_col = table_df_mapping[header_index]
            totals[total_header] = df_source[df_col].sum()

        for paragraph in doc.paragraphs:
            for total_header, total_value in totals.items():
                placeholder = f"{{{total_header}}}"
                if placeholder in paragraph.text:
                    words = num2words(int(total_value), lang='es').capitalize()
                    cents = int(round((total_value - int(total_value)) * 100))
                    formatted_value = f"{total_value:,.2f}".replace(",", " ")
                    replacement = f"${formatted_value} ({words} pesos {cents:02}/100)"
                    paragraph.text = paragraph.text.replace(placeholder, replacement)

        # Add rows for Subtotal, IVA, and Total
        for label in ["SUBTOTAL", "IVA", "GRAN TOTAL"]:
            row_cells = table.add_row().cells
            row_cells[0].text = label
            for total_header, total_value in totals.items():
                col_index = table_headers.index(total_header)
                if label == "IVA":
                    row_cells[col_index].text = "0.00"
                else:
                    row_cells[col_index].text = f"${total_value:,.2f}"

    return doc


def show_doc_tables(word_file):
    # Open the Word document
    doc = Document(word_file)
    
    # Iterate through all tables in the document
    for i, table in enumerate(doc.tables, start=1):
        print(f"Table {i}:")
        
        # Fetch the first three rows (or fewer if the table has fewer rows)
        rows_to_show = min(5, len(table.rows))
        for j in range(rows_to_show):
            # Get the cells in the row and join their text contents
            row_data = [cell.text.strip() for cell in table.rows[j].cells]
            print(f"  Row {j + 1}: {row_data}")
        
        print("-" * 40)  # Separator for clarity

def save_to_word(word_document, word_file):
    output_file = word_file.replace('.docx', '_updated.docx')
    try:
        word_document.save(output_file)
        print(f"Document saved successfully as {output_file}.")
    except PermissionError:
        print(f"Error: Unable to save. Please close {output_file} and try again.")

def save_as_pdf(word_file):
    """
    Save a Word document as a PDF, replacing fields with the first row from the Excel source.
    """
    # Check if the Word file exists
    if not os.path.exists(word_file):
        print(f"Error: The file '{word_file}' does not exist. Please check the path and try again.")
        return

    # Define the output PDF file path
    pdf_file = word_file.replace(".docx", ".pdf")
    print(f"Word file: {os.path.abspath(word_file)}")
    print(f"PDF file: {os.path.abspath(pdf_file)}")

    # Open Word application
    try:
        word = win32.gencache.EnsureDispatch("Word.Application")
    except Exception as e:
        print(f"Error: Failed to initialize Word application. {e}")
        return

    doc = None
    try:
        # Open the Word document
        doc = word.Documents.Open(os.path.abspath(word_file))

        # Export as PDF
        doc.ExportAsFixedFormat(
            OutputFileName=os.path.abspath(pdf_file),
            ExportFormat=17,  # PDF format
            OpenAfterExport=False,
            OptimizeFor=0,  # Print optimization
            CreateBookmarks=1  # Create bookmarks from headings
        )
        print(f"PDF successfully created: {os.path.abspath(pdf_file)}")
    except Exception as e:
        print(f"Error: Unable to save as PDF. {e}")
    finally:
        if doc:
            try:
                doc.Close(False)
            except Exception as close_error:
                print(f"Warning: Could not close the document properly: {close_error}")
        word.Quit()

# Generar el Excel de Precios Compranet
def generador_propuesta_economica_excel():
    xlsx_template = 'LA-12-NEF-012NEF001-I-1-2025_template_compranet.xlsx'
    source_file = './LA-12-NEF-012NEF001-I-1-2025 Base para PT y PE.xlsx'
    output_folder = './Output'

    if not os.path.exists(source_file):
        print("Source file does not exist.")
        return

    df_source = pd.read_excel(source_file, sheet_name='Core')
    if df_source.empty:
        print("Dataframe source is empty.")
        return

    print("Dataframe source found. Do we proceed extracting the dictionary? (yes/no)")
    while True:
        response = input().strip().lower()
        if response == 'yes':
            economic_data = extract_dictionary(source_file)
            break
        elif response == 'no':
            return
        else:
            print("Please answer 'yes' or 'no'.")

    if economic_data and os.path.exists(xlsx_template):
        print("Do we proceed generating xlsx to upload prices? (yes/no)")
        while True:
            response = input().strip().lower()
            if response == 'yes':
                write_economic_data(xlsx_template, economic_data, output_folder)
                break
            elif response == 'no':
                return
            else:
                print("Please answer 'yes' or 'no'.")
    else:
        print("Economic data is empty or template file is missing.")    



def normalize_string_case_insensitive(s):
    """Normalize string for case-insensitive comparison."""
    if isinstance(s, str):
        return ' '.join(s.strip().upper().split())  # Convert to uppercase for case-insensitive comparison
    return s

def write_economic_data(input_xlsx, dictionary, output_folder):
    header_row = 6
    headers = ['DESCRIPCION DETALLADA', 'PRECIO UNITARIO SIN IMPUESTOS', 'MONTO DE LA OFERTA SIN IMPUESTOS', 'IVA', 'OTROS IMPUESTOS', 'MONTO TOTAL DE LA OFERTA']

    # Load the Excel workbook
    workbook = openpyxl.load_workbook(input_xlsx)
    sheet = workbook.active

    # Extract headers from the header row
    excel_headers = [sheet.cell(row=header_row, column=col).value for col in range(1, sheet.max_column + 1)]

    # Check for missing headers
    missing_headers = [header for header in headers if header not in excel_headers]
    if missing_headers:
        print(f"Missing headers: {missing_headers}")
        return

    print("All expected headers are present.")

    # Normalize headers for indexing
    normalized_excel_headers = [normalize_string_case_insensitive(header) for header in excel_headers]

    # Locate and replace data
    not_found = []
    for desc, values in dictionary.items():
        normalized_desc = normalize_string_case_insensitive(desc)
        found = False
        for row in range(header_row + 1, sheet.max_row + 1):
            cell_value = sheet.cell(row=row, column=normalized_excel_headers.index(normalize_string_case_insensitive('DESCRIPCION DETALLADA')) + 1).value
            if normalize_string_case_insensitive(cell_value) == normalized_desc:
                found = True
                for key, value in values.items():
                    if key in headers:
                        col = normalized_excel_headers.index(normalize_string_case_insensitive(key)) + 1
                        sheet.cell(row=row, column=col, value=value)
                break
        if not found:
            not_found.append(desc)

    if not_found:
        print("Descriptions not found:")
        for desc in not_found:
            print(f"- {desc}")

    # Save the updated workbook
    os.makedirs(output_folder, exist_ok=True)
    output_path = os.path.join(output_folder, "P01 Precios Compranet.xlsx")
    workbook.save(output_path)
    print(f"File saved to {output_path}")

def extract_dictionary(xlsx_file):
    df = pd.read_excel(xlsx_file, sheet_name='Core')
    
    if 'Descripción' not in df.columns or 'Cantidad Máxima' not in df.columns or 'Precio Unitario' not in df.columns:
        print("Required columns are missing in the input file.")
        return {}

    dictionary = {}
    for _, row in df.iterrows():
        descripcion = row['Descripción']
        cantidad_maxima = row['Cantidad Máxima']
        precio_unitario = row['Precio Unitario']

        if isinstance(cantidad_maxima, (int, float)) and isinstance(precio_unitario, (int, float)) and cantidad_maxima > 0 and precio_unitario > 0:
            dictionary[descripcion] = {
                'DESCRIPCION DETALLADA': descripcion,
                'PRECIO UNITARIO SIN IMPUESTOS': precio_unitario,
                'MONTO DE LA OFERTA SIN IMPUESTOS': precio_unitario * cantidad_maxima,
                'IVA': 0,
                'OTROS IMPUESTOS': 0,
                'MONTO TOTAL DE LA OFERTA': precio_unitario * cantidad_maxima
            }
    
    return dictionary

# Generar la propuesta económica en excel 

def populate_excel():
    print("La función no está lista, abre manualmente el word y el excel y pega la propuesta económica en el template")

#Print Bookmarks
def printBookmarks(word_file):
    """
    Prints the first-level headings (bookmarks) in the Word document.
    
    Args:
        word_file (str): Path to the Word document.
    """
    doc = Document(word_file)
    header_count = 1  # Start counting from 1
    
    for paragraph in doc.paragraphs:
        if paragraph.style.name.startswith('Heading 1'):
            print(f"Header: {header_count}, Bookmark: {paragraph.text}")
            header_count += 1

# Orquestador

def main():
    word_file = os.path.join(working_folder, 'Cartas.docx')
    doc = Document(word_file)
    excel_file = os.path.join(working_folder, 'Cartas.xlsx')

    df_raw = pd.read_excel(excel_file, sheet_name='Core')
    df_raw = df_raw.sort_values(by='NUMERO DE PARTIDA')
    df_eseotres = df_raw[df_raw['Membrete'] == 'Eseotres']
    df_venus = df_raw[df_raw['Membrete'] == 'Venus']
    df_rafarm = df_raw[df_raw['Membrete'] == 'Rafarm']
    df_labopharma = df_raw[df_raw['CLAVE (12 DÍGITOS)'] == '010.000.4413.00']
    human_check_word = os.path.join(working_folder, 'Cartas_updated.docx')
    
    # Tabla de normas
    normas_table = 7
    normas_headers_row = 1
    normas_headers = ['PARTIDA', 'CLAVE DEL COMPENDIO NACIONAL DE INSUMOS PARA LA SALUD', 'DENOMINACIÓN', 'NORMA']
    normas_types = {
        'NUMERO DE PARTIDA': 'string',
        'CLAVE (12 DÍGITOS)': 'string',
        'Descripción': 'string',
        'Norma': 'string'
    }
    normas_total = False
    normas_df_mapping = ['NUMERO DE PARTIDA', 'CLAVE (12 DÍGITOS)', 'Descripción', 'Norma']
    normas_total_headers = []
    
    #Propuesta económica
    económica_table = 29
    económica_headers_row = 2
    económica_headers = ['No.\nPartida', 'Clave \n(12 dígitos)', 'Descripción del bien ofertado', 'Uni', 'Cant', 'Tipo', 'Mínimo', 'Máximo', 'Mínimo', 'Máximo', 'Unidad de Medida', 'Registro Sanitario', 'Fabricante del bien', 'Precio Unitario', 'Bianual Mínimo', 'Bianual Máximo']
    economica_df_mapping = ['NUMERO DE PARTIDA', 'CLAVE (12 DÍGITOS)', 'Descripción', 'UNI', 'CANT', 'TIPO', 'Cantidad Mínima', 'Cantidad Máxima', 'Cantidad Mínima', 'Cantidad Máxima', 'Unidad de Medida', 'NÚMERO DE REGISTRO SANITARIO','FABRICANTE', 'Precio Unitario', 'Importe total Mínimo', 'Importe total Máximo']
    economica_types = { 
        'NUMERO DE PARTIDA': 'string',
        'CLAVE (12 DÍGITOS)': 'string',
        'Descripción': 'string',
        'UNI': 'string',
        'CANT': 'string',
        'TIPO': 'string',
        'Cantidad Mínima': 'number',
        'Cantidad Máxima': 'number',
        'Cantidad Mínima': 'number',
        'Cantidad Máxima': 'number',
        'Unidad de Medida': 'string',
        'NÚMERO DE REGISTRO SANITARIO': 'string',
        'FABRICANTE': 'string',
        'Precio Unitario': 'currency',
        'Importe total Mínimo': 'currency',
        'Importe total Máximo': 'currency'
        }
    economica_total = True
    economica_total_headers = ['Bianual Mínimo', 'Bianual Máximo'] 
    
    tecnica_table = 31
    tecnica_headers_row = 2
    tecnica_headers = ['NUMERO DE PARTIDA', 'GPO', 'GEN', 'ESP', 'DF', 'NOMBRE GENÉRICO', 'Descripción', 'UNI', 'CANT', 'TIPO', 'Mínima', 'Máxima', 'MARCA O DENOMINACIÓN DISTINTIVA', 'FABRICANTE', 'PAÍS DE ORIGEN', 'NÚMERO DE REGISTRO SANITARIO', 'CÓDIGO DE BARRAS (CUANDO APLIQUE)']
    tecnica_df_mapping = ['NUMERO DE PARTIDA', 'GPO','GEN','ESP', 'DF', 'NOMBRE GENÉRICO','Descripción', 'UNI', 'CANT','TIPO','Cantidad Mínima',  'Cantidad Máxima', 'MARCA O DENOMINACIÓN DISTINTIVA', 'FABRICANTE', 'PAÍS DE ORIGEN','NÚMERO DE REGISTRO SANITARIO', 'CÓDIGO DE BARRAS']
    tecnica_types = { 
        'NUMERO DE PARTIDA': 'string',
        'GPO': 'string',
        'GEN': 'string',
        'ESP': 'string',
        'DF': 'string',
        'NOMBRE GENÉRICO': 'string',
        'Descripción': 'string',
        'UNI': 'string',
        'CANT': 'number',
        'TIPO': 'string', 
        'Cantidad Mínima': 'number', 
        'Cantidad Máxima': 'number',
        'MARCA O DENOMINACIÓN DISTINTIVA': 'string',
        'FABRICANTE': 'string',
        'PAÍS DE ORIGEN': 'string',
        'NÚMERO DE REGISTRO SANITARIO': 'string', 
        'CÓDIGO DE BARRAS': 'string'
        }    
    tecnica_total = False
    tecnica_total_headers = []

    #Tabla de normas 2
    normas_table_2 = 33
    
    #Tabla con marbete eseotres para todas las claves de Caducidad

    caducidadTodosFabricantes_row = 1
    caducidadTodosFabricantes_headers = ['PART NO.', 'CLAVE', 'DESCRIPCIÓN BREVE', 'CADUCIDAD MÍNIMA DE LOS BIENES']

    caducidadTodosFabricantes_df_mapping = ['NUMERO DE PARTIDA','CLAVE (12 DÍGITOS)', 'NOMBRE GENÉRICO', 'Caducidad mínima']
    caducidadTodosFabricantes_types = { 
        'NUMERO DE PARTIDA': 'string',
        'CLAVE (12 DÍGITOS)': 'string',
        'NOMBRE GENÉRICO': 'string',
        'Caducidad mínima': 'string'
    }
    caducidadTodosFabricantes_total = False
    caducidadTodosFabricantes_total_headers = []
    
    caducidadEseotres = 34
    caducidadRafarm = 35
    caducidadLabopharma = 36

    # Rafarm Acuerdo
    Rafarm_acuerdo_table = 41
    Rafarm_acuerdo_row = 1
    Rafarm_acuerdo_headers = ['CLAVE', 'DESCRIPCIÓN', 'CANTIDAD MÁXIMA REQUERIDA 2025-2026', 'REGISTRO SANITARIO']
    Rafarm_acuerdo_types = {
        'CLAVE': 'string',
        'DESCRIPCIÓN': 'string',
        'Cantidad Máxima': 'number',
        'REGISTRO SANITARIO': 'string'
    }
    Rafarm_acuerdo_df_mapping = ['CLAVE (12 DÍGITOS)', 'Descripción', 'Cantidad Máxima', 'NÚMERO DE REGISTRO SANITARIO']
    ##Rafarm apoyo
    Rafarm_apoyo_table = 42
    Rafarm_apoyo_row = 1
    Rafarm_apoyo_headers = ['CLAVE', 'DESCRIPCIÓN', 'CANTIDAD MÁXIMA REQUERIDA\n2025-2026', 'REGISTRO SANITARIO', 'CANTIDAD O PORCENTAJE QUE RESPALDA']
    Rafarm_apoyo_types = {
        'CLAVE': 'string',
        'DESCRIPCIÓN': 'string',
        'Cantidad Máxima': 'number',
        'REGISTRO SANITARIO': 'string',
        'CANTIDAD O PORCENTAJE QUE RESPALDA': 'string'
    }
    Rafarm_apoyo_df_mapping = ['CLAVE (12 DÍGITOS)', 'Descripción', 'Cantidad Máxima', 'NÚMERO DE REGISTRO SANITARIO','% RESPALDADO']
    Rafarm_total_headers = []
    Rafarm_total = False
    # Labopharm acuerdo
    LABATEC_acuerdo_table = 43
    LABATEC_acuerdo_row = 1
    LABATEC_acuerdo_headers = ['CLAVE', 'DESCRIPCIÓN', 'CANTIDAD MÁXIMA REQUERIDA 2025-2026', 'REGISTRO SANITARIO']
    LABATEC_acuerdo_types = {
        'CLAVE': 'string',
        'DESCRIPCIÓN': 'string',
        'Cantidad Máxima': 'number',
        'REGISTRO SANITARIO': 'string'
    }
    LABATEC_acuerdo_df_mapping = ['CLAVE (12 DÍGITOS)', 'Descripción', 'Cantidad Máxima', 'NÚMERO DE REGISTRO SANITARIO']
    ##Labopharma Apoyo
    LABATEC_apoyo_table = 45
    LABATEC_apoyo_row = 1
    LABATEC_apoyo_headers = ['CLAVE', 'DESCRIPCIÓN', 'CANTIDAD MÁXIMA REQUERIDA\n2025-2026', 'REGISTRO SANITARIO', 'CANTIDAD O PORCENTAJE QUE RESPALDA']
    LABATEC_apoyo_types = {
        'CLAVE': 'string',
        'DESCRIPCIÓN': 'string',
        'Cantidad Máxima': 'number',
        'REGISTRO SANITARIO': 'string',
        'CANTIDAD O PORCENTAJE QUE RESPALDA': 'string'
    }
    LABATEC_apoyo_df_mapping = ['CLAVE (12 DÍGITOS)', 'Descripción', 'Cantidad Máxima', 'NÚMERO DE REGISTRO SANITARIO','% RESPALDADO']
    LABATEC_total_headers = []
    LABATEC_total = False

    while True:
        print("\nMenu:")
        print("1) Show tables of the Word document: working")
        print("2) Populate tables: working")
        print("3) Save word to PDF preserving headers: working")
        print("4) Genera el excel de propuesta económica")
        print("5) Imprime los headers del archivo poblado")
        choice = input("Choose an option (1, 2, 3, 4, 5): ")
        
        if choice == '1':
            show_doc_tables(word_file)
        elif choice == '2':
            #Normas
            doc = populate_table(doc, df_raw, normas_table, normas_headers_row, normas_headers, normas_types, normas_total, normas_df_mapping, normas_total_headers)
            #Normas 2
            doc = populate_table(doc, df_raw, normas_table_2, normas_headers_row, normas_headers, normas_types, normas_total, normas_df_mapping, normas_total_headers)
            #Económica
            doc = populate_table(doc, df_raw, económica_table, económica_headers_row, económica_headers, economica_types, economica_total, economica_df_mapping, economica_total_headers)
            #Técncica
            doc = populate_table(doc, df_raw, tecnica_table, tecnica_headers_row, tecnica_headers, tecnica_types, tecnica_total, tecnica_df_mapping, tecnica_total_headers)
            #Caducidad SO3
            doc = populate_table(doc, df_raw, caducidadEseotres, caducidadTodosFabricantes_row, caducidadTodosFabricantes_headers, caducidadTodosFabricantes_types, caducidadTodosFabricantes_total, caducidadTodosFabricantes_df_mapping, caducidadTodosFabricantes_total_headers)
            #Caducidad RAFARM
            doc = populate_table(doc, df_raw, caducidadRafarm, caducidadTodosFabricantes_row, caducidadTodosFabricantes_headers, caducidadTodosFabricantes_types, caducidadTodosFabricantes_total, caducidadTodosFabricantes_df_mapping, caducidadTodosFabricantes_total_headers)
            #Caducidad Labopharma            
            doc = populate_table(doc, df_raw, caducidadLabopharma, caducidadTodosFabricantes_row, caducidadTodosFabricantes_headers, caducidadTodosFabricantes_types, caducidadTodosFabricantes_total, caducidadTodosFabricantes_df_mapping, caducidadTodosFabricantes_total_headers)
            #RAFARM acuerdo
            doc = populate_table(doc, df_raw, Rafarm_acuerdo_table, Rafarm_acuerdo_row, Rafarm_acuerdo_headers, Rafarm_acuerdo_types, Rafarm_total, Rafarm_acuerdo_df_mapping, Rafarm_total_headers)
            #RAFARM apoyo
            doc = populate_table(doc, df_raw, Rafarm_apoyo_table, Rafarm_apoyo_row, Rafarm_apoyo_headers, Rafarm_apoyo_types, Rafarm_total, Rafarm_apoyo_df_mapping, Rafarm_total_headers)
            save_to_word(doc, word_file)
            #LABATEC Acuerdo
            doc = populate_table(doc, df_raw, LABATEC_acuerdo_table, LABATEC_acuerdo_row, LABATEC_acuerdo_headers, LABATEC_acuerdo_types, LABATEC_total, LABATEC_acuerdo_df_mapping, LABATEC_total_headers)   
            #LABATEC Apoyo
            doc = populate_table(doc, df_raw, LABATEC_apoyo_table, LABATEC_apoyo_row, LABATEC_apoyo_headers, LABATEC_apoyo_types, LABATEC_total, LABATEC_apoyo_df_mapping, LABATEC_total_headers)

            save_to_word(doc, word_file)       
        elif choice == '3': 

            save_as_pdf(human_check_word)
        elif choice == '4': 
            generador_propuesta_economica_excel()
        elif choice == '5': 
            printBookmarks(human_check_word)
        else:
            print("Invalid choice. Please select either 1, 2, or 3.")
            continue  # Ask again if the input is invalid
        
        break  # Exit the loop if the user made a valid choice

if __name__ == "__main__":
    main()
